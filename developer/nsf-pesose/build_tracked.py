#!/usr/bin/env python3
"""
Apply tracked-change edits to Ruth's original abstract.

Reads:  developer/nsf-pesose/RM NSF ABSTRACT.docx   (pristine, untouched)
Writes: developer/nsf-pesose/RM NSF ABSTRACT (SG tracked).docx

Tracked insertions (w:ins, author "Stephen Gadd"):
  * new WP3 bullet on production back-end infrastructure for large-scale /
    commercial consumers (back-end/hosting, distinct from application code);
  * two new sections after WP4: "Security and Quality Control" (Criterion #5)
    and "Licensing" (Criterion #4).
Comments (native Word comments) anchored to the relevant phrase.

Insertions are real revisions: open in LibreOffice/Word with "Track Changes"
and accept/reject each one.
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(__file__).with_name("RM NSF ABSTRACT.docx")
OUT = Path(__file__).with_name("RM NSF ABSTRACT (SG tracked).docx")

AUTHOR = "Stephen Gadd"
DATE = "2026-07-21T09:00:00Z"
INITIALS = "SG"

_rev = [1000]


def _rid():
    _rev[0] += 1
    return str(_rev[0])


def _mark(el):
    el.set(qn("w:id"), _rid())
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), DATE)


# ------------------------------------------------------------------ tracked-insert helpers
def clone_blank(ref_p):
    """Deep-copy a reference <w:p>, drop its content, keep pPr (style/numbering),
    and mark its paragraph mark as an inserted revision."""
    p = deepcopy(ref_p)
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    ins = OxmlElement("w:ins")
    _mark(ins)
    rPr.insert(0, ins)  # ins/del come first in CT_ParaRPr
    return p


def add_ins_run(p, text, *, bold=False, underline=False, italic=False):
    ins = OxmlElement("w:ins")
    _mark(ins)
    r = OxmlElement("w:r")
    if bold or underline or italic:
        rpr = OxmlElement("w:rPr")
        if bold:
            rpr.append(OxmlElement("w:b"))
        if italic:
            rpr.append(OxmlElement("w:i"))
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rpr.append(u)
        r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    p.append(ins)
    return p


# ------------------------------------------------------------------ comment helper
def isolate_phrase(paragraph, phrase):
    """Split a paragraph's runs so `phrase` is its own run; return that run.
    Preserves the first run's character formatting; used only on plain body text."""
    runs = paragraph.runs
    full = "".join(r.text for r in runs)
    idx = full.find(phrase)
    if idx < 0:
        raise ValueError(f"phrase not found: {phrase!r} in {full[:60]!r}")
    template_rpr = None
    if runs:
        rpr = runs[0]._element.find(qn("w:rPr"))
        if rpr is not None:
            template_rpr = deepcopy(rpr)
    before, after = full[:idx], full[idx + len(phrase):]
    for r in runs:
        r._element.getparent().remove(r._element)

    def mk(txt):
        run = paragraph.add_run(txt)
        if template_rpr is not None:
            ex = run._element.find(qn("w:rPr"))
            if ex is not None:
                run._element.remove(ex)
            run._element.insert(0, deepcopy(template_rpr))
        return run

    if before:
        mk(before)
    target = mk(phrase)
    if after:
        mk(after)
    return target


# =====================================================================
doc = Document(str(SRC))
P = doc.paragraphs

# --- capture reference elements up front (indices are valid at load time) ---
p_scale = P[15]      # Building-on-a-Prototype body ("...50,484,667...")
p_gadd = P[37]       # WP2 bullet a
p_wp3_intro = P[45]  # WP3 intro
p_graph = P[46]      # WP3 bullet a
p_api = P[49]        # WP3 bullet d
p_sm_sec = P[64]     # "Still Missing" > Security/quality-control bullet
clone_normal = P[44]._p   # a blank Normal paragraph (clean clone source)
wp3_last_bullet = P[51]._p
before_still_missing = P[60]._p  # blank line before "Still Missing"

SECURITY_TEXT = (
    "Because the ecosystem invites external contributions and depends on an "
    "AI-assisted extraction pipeline, security and quality control are built into "
    "its design rather than treated as an afterthought. Access is authenticated "
    "through per-user API keys and tokens issued via the contributor profile, "
    "with capped, metered usage that both protects the service and underpins the "
    "commercial API tier described in Work Packages 1 and 3. Contribution "
    "requires a verified ORCiD identity, and datasets are reviewed by the "
    "editorial board before publication, with each record carrying its source and "
    "licence provenance. Per-contributor and per-contribution rollback allows any "
    "submission to be cleanly reverted, and changes to the index are versioned "
    "and auditable. Place attestations produced by the extraction pipeline (Work "
    "Package 2) are generated with recorded source-to-attestation lineage, "
    "versioned models, and human review before entering the shared index. The "
    "codebase is maintained with dependency and vulnerability scanning, code "
    "review, and controlled, signed releases, and the index and its history are "
    "preserved through redundant, distributed storage on LOCKSS-type principles "
    "together with dark archiving. Together these measures address the "
    "“secure” mandate of the PESOSE program across identity, integrity, "
    "supply chain, and preservation."
)

LICENSING_TEXT = (
    "WHG code, data, and curriculum are openly licensed, and all software and "
    "derived outputs produced under this project will be released under open "
    "licences. The index aggregates content from many independent sources under a "
    "range of terms: much is fully open (CC0 or CC-BY), while some carries "
    "share-alike (for example, GB1900 is CC-BY-SA) or non-commercial conditions. "
    "Our approach records each source's licence as provenance metadata carried "
    "with every record, surfaces attribution at the platform level, and honours "
    "share-alike and non-commercial terms in any commercial API tier rather than "
    "relicensing restricted content. This allows the ecosystem to generate "
    "revenue from permissively licensed content and value-added services while "
    "fully respecting the rights of upstream contributors, and gives downstream "
    "users clear, machine-readable licence information for every place record "
    "they consume."
)

BACKEND_BULLET = (
    "Provision and operate the production back-end infrastructure — elastic "
    "compute, storage and indexing capacity, redundancy, and monitored uptime "
    "— required to serve an index of hundreds of millions of records to "
    "large-scale and commercial consumers reliably. This hosting and operations "
    "capacity is distinct from the application code and is a recurring cost that "
    "the revenue model in Work Package 1 is designed to help sustain."
)

# ------------------------------------------------------------------ comments
doc.add_comment(
    [isolate_phrase(p_scale, "50,484,667 named entities drawn from 22 source gazetteers")],
    text=(
        "These ~50M entities are mostly places, and new sources mostly add "
        "attestations/toponyms of places already indexed rather than new entities, "
        "so entity count grows slowly. “Hundreds of millions” is accurate "
        "for attestations / place references, not for entities — worth aligning "
        "that usage (cf. “hundreds of millions of place references” in the "
        "Workplan vs “place entities” in WP3) so it doesn't read as an "
        "order-of-magnitude jump in entities."
    ),
    author=AUTHOR, initials=INITIALS,
)

doc.add_comment(
    [isolate_phrase(p_gadd, "(Gadd, 2026)")],
    text=(
        "This refers to the live GOTW proof-of-concept (acknowledged as such, "
        "surfacing the achievable refinements). GB-STAMP is a separate experiment, "
        "to be published later as a processing exemplar and dataset — so don't "
        "cite it as complete. Worth a phrase noting extracted attestations pass "
        "human review and carry recorded provenance before entering the shared "
        "index (Criterion #5)."
    ),
    author=AUTHOR, initials=INITIALS,
)

doc.add_comment(
    [isolate_phrase(p_wp3_intro, "build on with confidence")],
    text=(
        "This is where the large-scale / commercial-consumer point lands: serving "
        "hundreds of millions of records to those users reliably is a back-end "
        "problem (compute, storage, indexing capacity, redundancy, monitored "
        "uptime), not just a software one. I've inserted a WP3 bullet making that "
        "explicit and tying the recurring hosting cost to the WP1 revenue model "
        "— those consumers are exactly who that model charges. It also "
        "strengthens the funding case: robust infrastructure, not only code, is "
        "what the investment buys."
    ),
    author=AUTHOR, initials=INITIALS,
)

doc.add_comment(
    [isolate_phrase(p_graph, "an open-source graph platform")],
    text=(
        "This is already scoped, so we can state it with confidence rather than as "
        "an aspiration: the v4 graph/attestation model is documented "
        "(docs.whgazetteer.org/content/v4.html) and built on the Pelagios "
        "place-attestation-ontology, which ISHI/WHG is leading. Suggest citing both "
        "to preempt the “which platform / is it designed yet?” question."
    ),
    author=AUTHOR, initials=INITIALS,
)

doc.add_comment(
    [isolate_phrase(p_api, "usage metering for commercial users")],
    text=(
        "Part of this is already live: API keys/tokens are issued via the User "
        "Profile and usage is capped and metered, so the commercial-metering story "
        "is real today (worth saying so; it also cross-references WP1's revenue "
        "model). The genuinely new items are real-time feeds and webhooks. That "
        "existing key + metering layer is also most of the Criterion #5 "
        "identity/access-management answer."
    ),
    author=AUTHOR, initials=INITIALS,
)

doc.add_comment(
    [isolate_phrase(p_sm_sec, "Security/quality-control Infrastructure (Criterion #5)")],
    text=(
        "Criteria #4 and #5 are now drafted below as inserted sections "
        "(“Security and Quality Control” and “Licensing”). Once "
        "accepted, these two “Still Missing” bullets can be removed."
    ),
    author=AUTHOR, initials=INITIALS,
)

# ------------------------------------------------------------------ tracked insertions
# 1) new WP3 back-end bullet, right after the last WP3 bullet
beb = clone_blank(wp3_last_bullet)
add_ins_run(beb, BACKEND_BULLET)
wp3_last_bullet.addnext(beb)

# 2) two new sections after WP4 (before the blank preceding "Still Missing")
seq = []
seq.append(clone_blank(clone_normal))  # spacer
h1 = clone_blank(clone_normal); add_ins_run(h1, "Security and Quality Control", bold=True, underline=True); seq.append(h1)
b1 = clone_blank(clone_normal); add_ins_run(b1, SECURITY_TEXT); seq.append(b1)
seq.append(clone_blank(clone_normal))  # spacer
h2 = clone_blank(clone_normal); add_ins_run(h2, "Licensing", bold=True, underline=True); seq.append(h2)
b2 = clone_blank(clone_normal); add_ins_run(b2, LICENSING_TEXT); seq.append(b2)
for el in seq:
    before_still_missing.addprevious(el)

doc.save(str(OUT))
print("wrote", OUT)
